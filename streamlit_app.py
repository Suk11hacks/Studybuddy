# streamlit_app.py
import streamlit as st
import google.generativeai as genai
import os
import fitz  # PyMuPDF for PDF parsing
import tempfile
import whisper  # For speech-to-text
from pytube import YouTube  # For YouTube video downloads
from docx import Document
from fpdf import FPDF
import sqlite3
from datetime import datetime
import requests
from bs4 import BeautifulSoup
from PIL import Image

# Load API key securely from Streamlit secrets
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=GEMINI_API_KEY)

# Initialize Gemini models
text_model = genai.GenerativeModel("gemini-pro")
vision_model = genai.GenerativeModel("gemini-pro-vision")

# Initialize SQLite database for session tracking
conn = sqlite3.connect("user_sessions.db", check_same_thread=False)
cursor = conn.cursor()
cursor.execute("""
CREATE TABLE IF NOT EXISTS sessions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    timestamp TEXT,
    user_query TEXT,
    notes TEXT,
    summary TEXT,
    flashcards TEXT
)
""")
conn.commit()

# UI
st.title("🔬 StudyBuddy")
st.markdown("Generate detailed notes from MCQs, PDFs, images, and video lectures using web search + Gemini.")

query = st.text_input("Enter your MCQ question or topic:")
additional_instructions = st.text_area("Add custom instructions for note generation:")
note_formatting = st.text_input("Specify formatting style (e.g., bullet points, tables, sections):")
youtube_url = st.text_input("Paste a YouTube video URL (optional):")
show_summary = st.checkbox("Show summary of content before generating notes")

uploaded_pdf = st.file_uploader("Upload a PDF (up to 500 pages):", type=["pdf"])
uploaded_video = st.file_uploader("Upload a video lecture (optional):", type=["mp4", "mov", "avi"])
uploaded_image = st.file_uploader("📷 Upload an image of a chemistry question:", type=["png", "jpg", "jpeg"])

# Flashcard and quiz material section
st.markdown("### 🧠 Revision Material Generator")
generate_flashcards = st.checkbox("Generate flashcards and MCQs from the notes")

# Extract text from uploaded PDF
pdf_context = ""
if uploaded_pdf is not None:
    try:
        with fitz.open(stream=uploaded_pdf.read(), filetype="pdf") as doc:
            for page in doc:
                pdf_context += page.get_text()
    except Exception as e:
        st.error(f"Error reading PDF: {e}")

# Transcribe uploaded or YouTube video audio to text
video_transcript = ""
if uploaded_video is not None:
    st.warning("Uploaded video transcription is currently disabled due to environment limitations.")
elif youtube_url:
    st.info("Processing YouTube video...")
    try:
        yt = YouTube(youtube_url)
        stream = yt.streams.filter(only_audio=True).first()
        audio_file = stream.download(filename="yt_audio.mp4")
        model_whisper = whisper.load_model("base")
        result = model_whisper.transcribe(audio_file)
        video_transcript += result["text"]
    except Exception as e:
        st.error(f"Error processing YouTube video: {e}")

# Web scraping function using Bing search and BeautifulSoup
def search_bing_scrape(query, max_results=5):
    headers = {"User-Agent": "Mozilla/5.0"}
    results = []
    try:
        resp = requests.get(f"https://www.bing.com/search?q={query}", headers=headers)
        soup = BeautifulSoup(resp.text, "html.parser")
        for item in soup.select(".b_algo")[:max_results]:
            title_elem = item.find("h2")
            snippet_elem = item.find("p")
            link_elem = item.find("a", href=True)
            if title_elem and link_elem:
                results.append({
                    "title": title_elem.text.strip(),
                    "snippet": snippet_elem.text.strip() if snippet_elem else "",
                    "url": link_elem["href"]
                })
    except Exception as e:
        st.error(f"Web scraping search failed: {e}")
        results.append({"title": "Search failed", "snippet": str(e), "url": "#"})
    return results

# Function to generate notes from context and question using Gemini text model
def generate_notes_with_context(question, context_chunks, pdf_text="", video_text="", custom="", fmt=""):
    full_context = "\n\n".join(context_chunks)
    prompt = f"""
You are a chemistry expert. Given the following question and context, generate detailed, structured notes in the following format:

Format: {fmt}

Question: {question}

User Instructions:
{custom}

Context from search:
{full_context}

Context from uploaded PDF:
{pdf_text}

Context from video lecture transcript:
{video_text}

If the question involves a test, include:
- Discovery
- Reagents and their interactions
- Observations and principles

If it's a mechanism, explain every step clearly.
"""
    response = text_model.generate_content(prompt)
    return response.text

# Main logic on button click
if st.button("Generate Notes"):

    # Priority: Image input first
    if uploaded_image:
        st.info("Processing image with Gemini Vision...")
        image = Image.open(uploaded_image)

        try:
            img_response = vision_model.generate_content(
                [
                    "You are a chemistry expert. Analyze this image of a question and generate detailed, structured notes explaining the content.",
                    image
                ]
            )
            notes = img_response.text
            st.subheader("🖼️ Notes from Image")
            st.write(notes)

            flashcard_output = ""
            if generate_flashcards:
                st.subheader("🧩 Flashcards and Practice Questions (from image notes)")
                flashcard_prompt = f"""
Based on the following chemistry notes extracted from an image, generate:
1. 5 multiple choice questions with 4 options each (indicate the correct one).
2. 5 flashcards (front with question/hint, back with explanation).

Notes:
{notes}
"""
                flashcard_output = text_model.generate_content(flashcard_prompt).text
                st.text_area("Generated Flashcards and MCQs:", flashcard_output, height=300)

            timestamp = datetime.now().isoformat()
            cursor.execute("""
            INSERT INTO sessions (timestamp, user_query, notes, summary, flashcards)
            VALUES (?, ?, ?, ?, ?)
            """, (timestamp, "Image Input", notes, "", flashcard_output if generate_flashcards else ""))
            conn.commit()

        except Exception as e:
            st.error(f"Gemini Vision processing failed: {e}")

    elif query:
        st.info("Searching web and generating notes from query...")
        search_results = search_bing_scrape(query)
        context_chunks = [res['snippet'] for res in search_results]

        notes = generate_notes_with_context(
            query,
            context_chunks,
            pdf_text=pdf_context,
            video_text=video_transcript,
            custom=additional_instructions,
            fmt=note_formatting
        )

        summary = ""
        if show_summary:
            summary_prompt = f"Summarize this content for a chemistry student:\n\n{pdf_context}\n\n{video_transcript}"
            summary = text_model.generate_content(summary_prompt).text
            st.subheader("📌 Summary")
            st.write(summary)

        st.subheader("🧪 Generated Notes")
        st.write(notes)

        with st.expander("🔗 Sources"):
            for res in search_results:
                st.markdown(f"- [{res['title']}]({res['url']})")

        flashcard_output = ""
        if generate_flashcards:
            st.subheader("🧩 Flashcards and Practice Questions")
            flashcard_prompt = f"""
Based on the following chemistry notes, generate:
1. 5 multiple choice questions with 4 options each (indicate the correct one).
2. 5 flashcards (front with question/hint, back with explanation).

Notes:
{notes}
"""
            flashcard_output = text_model.generate_content(flashcard_prompt).text
            st.text_area("Generated Flashcards and MCQs:", flashcard_output, height=300)

        timestamp = datetime.now().isoformat()
        cursor.execute("""
        INSERT INTO sessions (timestamp, user_query, notes, summary, flashcards)
        VALUES (?, ?, ?, ?, ?)
        """, (timestamp, query, notes, summary if show_summary else "", flashcard_output if generate_flashcards else ""))
        conn.commit()

    else:
        st.warning("Please enter a query or upload an image to generate notes.")

# View past sessions
with st.expander("🗂 View Saved Sessions"):
    cursor.execute("SELECT timestamp, user_query, notes FROM sessions ORDER BY timestamp DESC LIMIT 10")
    rows = cursor.fetchall()
    for row in rows:
        st.markdown(f"**Time:** {row[0]}")
        st.markdown(f"**Query:** {row[1]}")
        st.markdown(f"**Notes Preview:** {row[2][:300]}...")
        st.markdown("---")

# Export options outside button, showing last notes only if generated
if 'notes' in locals():
    # Export to PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in notes.split('\n'):
        pdf.multi_cell(0, 10, line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        pdf.output(f.name)
        st.download_button("📄 Download Notes as PDF", data=open(f.name, "rb").read(), file_name="chem_notes.pdf")

    # Export to DOCX
    doc = Document()
    doc.add_heading("AI-Generated Chemistry Notes", 0)
    for para in notes.split("\n"):
        doc.add_paragraph(para)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        doc.save(f.name)
        st.download_button("📝 Download Notes as DOCX", data=open(f.name, "rb").read(), file_name="chem_notes.docx")
